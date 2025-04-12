from docx import Document
from fpdf import FPDF

# Informações para o currículo
name = "Lúcio Motta"
title = "Back-End Developer"
contact = {
    "Email": "dev.lucio0@gmail.com",
    "Telefone": "(61) 9XXXX-XXXX",
    "LinkedIn": "linkedin.com/in/lucio-motta-828613166",
    "GitHub": "github.com/lucionmotta",
    "Portfólio": "portifolio-lucio-motta.netlify.app"
}
resumo = (
    "Desenvolvedor Back-End com foco em SvelteKit e ReactJS. "
    "Atualmente estagiando na Polícia do Senado Federal com experiência prática em Oracle APEX, "
    "integração com APIs e manipulação de dados em bancos Oracle. "
    "Apaixonado por soluções seguras, escaláveis e orientadas a dados. "
    "Cria projetos com impacto real, aplicando boas práticas em desenvolvimento."
)
habilidades = [
    "JavaScript, TypeScript, PHP, Python, PL/SQL",
    "ReactJS, SvelteKit, Express.js, Streamlit",
    "Oracle, PostgreSQL, MySQL",
    "Git, Docker, Oracle JET, Power BI",
    "REST APIs, Web Scraping, Machine Learning (scikit-learn), EmailJS"
]
experiencia = [
    {
        "titulo": "Estagiário em Desenvolvimento de Sistemas",
        "empresa": "Polícia do Senado Federal",
        "periodo": "2024 - Atual",
        "descricao": [
            "Desenvolvimento de sistemas internos com foco em segurança e confidencialidade.",
            "Utilização prática de Oracle APEX e banco de dados Oracle.",
            "Atuação com APIs e scripts em PL/SQL.",
            "Vivência com metodologias ágeis."
        ]
    }
]
educacao = "Análise e Desenvolvimento de Sistemas — Instituição Estácio, DF (Em andamento)"

projetos = [
    ["TopBolões", "Plataforma de apostas esportivas completa, feita do zero em PHP."],
    ["Treino Turbo App", "Aplicação ReactJS para controle de treinos de academia. (GitHub)"],
    ["Previsão de Preço de Pizza (ML)", "Machine Learning com Python para previsão de valores. (GitHub)"],
    ["Oral Care Dashboard", "Streamlit + Python para análise de dados odontológicos. (GitHub)"],
    ["Site Pizzaria", "Sistema completo com PHP e MySQL (GitHub)"]
]

# Criar o documento Word
doc = Document()
doc.add_heading(name, 0)
doc.add_paragraph(title, style='Intense Quote')

# Contato
doc.add_heading("Contato", level=1)
for key, value in contact.items():
    doc.add_paragraph(f"{key}: {value}")

# Resumo
doc.add_heading("Resumo Profissional", level=1)
doc.add_paragraph(resumo)

# Habilidades Técnicas
doc.add_heading("Habilidades Técnicas", level=1)
for item in habilidades:
    doc.add_paragraph(f"• {item}")

# Experiência Profissional
doc.add_heading("Experiência Profissional", level=1)
for exp in experiencia:
    doc.add_paragraph(exp["titulo"], style="Heading 2")
    doc.add_paragraph(f"{exp['empresa']} — {exp['periodo']}")
    for desc in exp["descricao"]:
        doc.add_paragraph(f"• {desc}")

# Educação
doc.add_heading("Educação", level=1)
doc.add_paragraph(educacao)

# Projetos em Destaque
doc.add_heading("Projetos em Destaque", level=1)
for proj in projetos:
    doc.add_paragraph(f"{proj[0]}: {proj[1]}")

# Salvar como .docx
word_path = "docs/Curriculo_Lucio_Motta.docx"
doc.save(word_path)

# Criar o PDF usando fpdf2
pdf = FPDF()
pdf.add_page()

# Adicionar uma fonte TrueType que suporte Unicode
pdf.add_font('DejaVu', '', 'fonts/DejaVuSans.ttf', uni=True)
pdf.set_font('DejaVu', '', 12)

# Cabeçalho
pdf.set_font("DejaVu", 'B', 16)
pdf.cell(0, 10, name, ln=True)
pdf.set_font("DejaVu", '', 12)
pdf.cell(0, 10, title, ln=True)
pdf.ln(5)

# Contato
pdf.set_font("DejaVu", 'B', 14)
pdf.cell(0, 10, "Contato", ln=True)
pdf.set_font("DejaVu", '', 12)
for key, value in contact.items():
    pdf.cell(0, 10, f"{key}: {value}", ln=True)
pdf.ln(5)

# Resumo
pdf.set_font("DejaVu", 'B', 14)
pdf.cell(0, 10, "Resumo Profissional", ln=True)
pdf.set_font("DejaVu", '', 12)
pdf.multi_cell(0, 10, resumo)
pdf.ln(5)

# Habilidades
pdf.set_font("DejaVu", 'B', 14)
pdf.cell(0, 10, "Habilidades Técnicas", ln=True)
pdf.set_font("DejaVu", '', 12)
for item in habilidades:
    pdf.cell(0, 10, f"• {item}", ln=True)
pdf.ln(5)

# Experiência
pdf.set_font("DejaVu", 'B', 14)
pdf.cell(0, 10, "Experiência Profissional", ln=True)
pdf.set_font("DejaVu", '', 12)
for exp in experiencia:
    pdf.set_font("DejaVu", 'B', 12)
    pdf.cell(0, 10, exp["titulo"], ln=True)
    pdf.set_font("DejaVu", '', 12)
    pdf.cell(0, 10, f"{exp['empresa']} — {exp['periodo']}", ln=True)
    for desc in exp["descricao"]:
        pdf.cell(0, 10, f"• {desc}", ln=True)
pdf.ln(5)

# Educação
pdf.set_font("DejaVu", 'B', 14)
pdf.cell(0, 10, "Educação", ln=True)
pdf.set_font("DejaVu", '', 12)
pdf.multi_cell(0, 10, educacao)
pdf.ln(5)

# Projetos
pdf.set_font("DejaVu", 'B', 14)
pdf.cell(0, 10, "Projetos em Destaque", ln=True)
pdf.set_font("DejaVu", '', 12)
for proj in projetos:
    pdf.multi_cell(0, 10, f"{proj[0]}: {proj[1]}")

# Salvar PDF
pdf_path = "docs/Curriculo_Lucio_Motta.pdf"
pdf.output(pdf_path)

print(f"Arquivos gerados:\nWord: {word_path}\nPDF: {pdf_path}")